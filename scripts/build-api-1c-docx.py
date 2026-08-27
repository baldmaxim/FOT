#!/usr/bin/env python
"""
Собирает Word- и Markdown-версию инструкции для 1С из API_1C.md.

Зачем скрипт, а не разовый файл: docx неизбежно устареет после правок Markdown,
а руками переносить 400 строк никто не будет. Инструкция правится в одном месте —
в API_1C.md, документ пересобирается командой.

Запуск (из корня репозитория):
    python scripts/build-api-1c-docx.py

Результат:
    docs/Инструкция_1С_обмен_табелями.docx  — для человека
    docs/Обмен_табелями_с_1С.md             — для агентов и внешних исполнителей

Берётся раздел «Закрытые согласованные табели» плюс вводная часть про адрес и
авторизацию: разработчику 1С нужен именно обмен табелями, а не generic-таблицы.
"""
from __future__ import annotations

import io
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm

ROOT = Path(__file__).resolve().parent.parent
SOURCE = ROOT / 'API_1C.md'
TARGET = ROOT / 'docs' / 'Инструкция_1С_обмен_табелями.docx'
# Та же выдержка в Markdown. Нужна там, где Word неудобен: агенты и внешние
# исполнители читают md заметно надёжнее. Собирается здесь же, чтобы не разошлась
# с docx и с исходником — ровно по той причине, что описана в шапке файла.
TARGET_MD = ROOT / 'docs' / 'Обмен_табелями_с_1С.md'

# Раздел, с которого начинается обмен табелями.
SECTION_MARKER = '# Закрытые согласованные табели — регулярный обмен с 1С'

INTRO = [
    ('heading', 1, 'Адрес и авторизация'),
    ('para', 'Базовый адрес: https://fot.su10.ru/api/public/v1'),
    ('para', 'В каждый запрос добавляется заголовок:'),
    ('code', 'Authorization: Bearer fot_<prefix>_<secret>'),
    ('para',
     'Токен выдаёт администратор ФОТ отдельно от этой инструкции. Он показывается '
     'один раз при создании ключа: если потерян, создаётся новый, восстановить '
     'существующий нельзя.'),
    ('para',
     'Машиночитаемый контракт всех методов — файл openapi-1c.yaml, он передаётся '
     'вместе с этим документом.'),
]


def add_code(doc: Document, text: str) -> None:
    """Блок кода: моноширинный шрифт на сером фоне абзаца."""
    for line in text.rstrip().split('\n'):
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.5)
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.space_before = Pt(0)
        run = para.add_run(line if line else ' ')
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x37)


def add_rich_text(para, text: str) -> None:
    r"""Инлайн-разметка: **жирный** и `код`. Остальное — как есть.

    ДВА ОГРАНИЧЕНИЯ, о которые легко споткнуться при правке API_1C.md:
      * разметка не переносится на следующую строку — каждая строка Markdown
        становится отдельным абзацем, и `**` из разорванной пары уедет в текст;
      * вложенность не разбирается: `**\`код\`**` отрендерится с кавычками.
    """
    for chunk in re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text):
        if not chunk:
            continue
        if chunk.startswith('**') and chunk.endswith('**'):
            para.add_run(chunk[2:-2]).bold = True
        elif chunk.startswith('`') and chunk.endswith('`'):
            run = para.add_run(chunk[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
        else:
            para.add_run(chunk)


def warn_broken_markup(line: str) -> None:
    """Ловит жирный, разорванный переносом строки.

    Каждая строка Markdown становится отдельным абзацем Word, поэтому пара `**`
    обязана открыться и закрыться в пределах одной строки. Иначе звёздочки уедут
    в текст документа, который уходит внешнему разработчику. Молчать про это
    нельзя — в готовом docx такое замечаешь последним.
    """
    if line.count('**') % 2:
        print('ВНИМАНИЕ: непарные ** — жирный разорван переносом строки:', file=sys.stderr)
        print(f'  {line}', file=sys.stderr)


def flush_table(doc: Document, rows: list[list[str]]) -> None:
    """Markdown-таблица → таблица Word со стилем сетки."""
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            if c >= len(table.columns):
                continue
            cell = table.cell(r, c)
            cell.text = ''
            para = cell.paragraphs[0]
            add_rich_text(para, cell_text)
            if r == 0:
                for run in para.runs:
                    run.bold = True
    doc.add_paragraph()


def build(md_text: str) -> Document:
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    doc.add_heading('ФОТ → 1С: обмен закрытыми табелями', level=0)
    subtitle = doc.add_paragraph(
        'Инструкция для разработчика 1С. Актуальна для методов /api/public/v1/timesheets.'
    )
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in subtitle.runs:
        run.italic = True

    for kind, *rest in INTRO:
        if kind == 'heading':
            doc.add_heading(rest[1], level=rest[0])
        elif kind == 'code':
            add_code(doc, rest[0])
        else:
            add_rich_text(doc.add_paragraph(), rest[0])

    lines = md_text.split('\n')
    start = next((i for i, l in enumerate(lines) if l.startswith(SECTION_MARKER)), None)
    if start is None:
        raise SystemExit(f'В {SOURCE.name} не найден раздел: {SECTION_MARKER}')

    in_code = False
    code_buffer: list[str] = []
    table_buffer: list[list[str]] = []

    for line in lines[start + 1:]:
        # Кодовые блоки: внутри них разметка не разбирается.
        if line.startswith('```'):
            if in_code:
                add_code(doc, '\n'.join(code_buffer))
                code_buffer = []
            in_code = not in_code
            continue
        if in_code:
            code_buffer.append(line)
            continue

        is_table_row = line.startswith('|') and line.endswith('|')
        if is_table_row:
            cells = [c.strip() for c in line.strip('|').split('|')]
            # Строка-разделитель (|---|---|) в документ не переносится.
            if all(set(c) <= set('-: ') for c in cells):
                continue
            table_buffer.append(cells)
            continue
        if table_buffer:
            flush_table(doc, table_buffer)
            table_buffer = []

        stripped = line.strip()
        if not stripped or stripped == '---':
            continue

        if stripped.startswith('#'):
            level = len(stripped) - len(stripped.lstrip('#'))
            title = stripped.lstrip('#').strip()
            # Нумерация разделов из Markdown в Word не нужна — там своя структура.
            title = re.sub(r'^\d+\.\s*', '', title)
            doc.add_heading(title.replace('`', ''), level=min(level, 4))
        elif stripped.startswith(('- ', '* ')):
            para = doc.add_paragraph(style='List Bullet')
            add_rich_text(para, stripped[2:])
        elif re.match(r'^\d+\.\s', stripped):
            para = doc.add_paragraph(style='List Number')
            add_rich_text(para, re.sub(r'^\d+\.\s', '', stripped))
        else:
            warn_broken_markup(stripped)
            add_rich_text(doc.add_paragraph(), stripped)

    if table_buffer:
        flush_table(doc, table_buffer)

    return doc


MD_HEADER = """<!--
  СОБИРАЕТСЯ АВТОМАТИЧЕСКИ из API_1C.md — руками не править.
  Пересборка: python scripts/build-api-1c-docx.py

  Здесь только раздел про обмен закрытыми табелями. Первая половина исходного
  документа описывает ДРУГОЙ API («сырые» таблицы на базовом адресе /external/v1)
  и к обмену табелями отношения не имеет — вырезана намеренно, чтобы не путать
  базовые адреса.
-->

**Базовый адрес:** `https://fot.su10.ru/api/public/v1`
**Авторизация:** заголовок `Authorization: Bearer <токен>`

---

"""


def extract_markdown(md_text: str) -> str:
    """Тот же срез, что уходит в Word: от маркера раздела до конца файла."""
    lines = md_text.splitlines()
    start = next((i for i, l in enumerate(lines) if l.startswith(SECTION_MARKER)), None)
    if start is None:
        raise SystemExit(f'В {SOURCE.name} не найден раздел: {SECTION_MARKER}')
    return MD_HEADER + '\n'.join(lines[start:]).rstrip() + '\n'


def main() -> None:
    if not SOURCE.exists():
        raise SystemExit(f'Не найден источник: {SOURCE}')
    md_text = io.open(SOURCE, encoding='utf-8').read()
    TARGET.parent.mkdir(parents=True, exist_ok=True)

    doc = build(md_text)
    doc.save(TARGET)
    print(f'Готово: {TARGET.relative_to(ROOT)}')

    io.open(TARGET_MD, 'w', encoding='utf-8', newline='').write(extract_markdown(md_text))
    print(f'Готово: {TARGET_MD.relative_to(ROOT)}')


if __name__ == '__main__':
    sys.exit(main())
